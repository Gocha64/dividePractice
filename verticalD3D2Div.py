import os, sys
sys.path.append(os.path.dirname(os.path.abspath(os.path.dirname(__file__))))
from verticalDigitNCalc import CalcPractice
from d3D2Div import d3D2Div
import docx
import verticalAnswerForD3D2Div


class verticalD3D2Div(CalcPractice):
    fileName = "d3D2Div"

    def __init__(self, baseForPracticeDocxName, baseForAnswerDocxName):
        CalcPractice.__init__(self, baseForPracticeDocxName, baseForAnswerDocxName)

    def getProblem(self):
        return d3D2Div()

    def fileWrite(self, file, c):
        file.write('\n' + c + '\n')

    def docxWrite(self, wFile, c):
        wFile.add_paragraph(c,style='나눗셈')
        wFile.add_paragraph('',style='나눗셈')
        wFile.add_paragraph('',style='나눗셈')

    def answering(self):
        verticalAnswerForD3D2Div.answering(self.fileName)

    def txt2Docx(self, wName):
        verticalAnswerForD3D2Div.answerText2Docx(self.fileName, wName)


if __name__=='__main__':
    asdf = verticalD3D2Div('baseForD3D2Div.docx', 'baseForVerticalAnswer.docx')
    print(asdf.fileName + " generator")
    i = int(input('생성 할 문제 수(1page = 20문제): '))
    asdf.digitNCalc(i)

'''
file = open("두자리 ÷ 세자리 산수.txt",'w')
print("기초 산수 5.0")
n = int(input("생성 할 문제 수: "))
i = 0
bank = [] #이전문제 저장
c = '' #문제

def check(c): #중복문제 체크
    for b in bank:
        if b == c:
            return True
    return False


while i < n:

    c = tenToHunDiv()
    if check(c):
        continue
    
    bank.append(c)

    if len(bank) > 51: #오래된 문제 삭제
        bank = bank[1:]

    file.write(c)
    i += 1

print("문제가 생성되었습니다.")
file.close()
'''
